import os
from typing import List
import docx

def read_txt(file_path: str) -> str:
    """Reads content from a text file."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def write_txt(file_path: str, content: str) -> None:
    """Writes content to a text file."""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)

def read_docx(file_path: str) -> str:
    """Reads text content from a Word document."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def write_docx(file_path: str, content: str) -> None:
    """Writes content to a Word document."""
    print(content)
    doc = docx.Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)
